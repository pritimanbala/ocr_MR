"""Word document parser backed by python-docx for DOCX files."""
from __future__ import annotations
from pathlib import Path
from typing import Any
from zipfile import ZipFile
from engineering_di.models import BoundingBox, Document, Image, Page, Table, TableCell, TextBlock, Token
from engineering_di.parsers.base_parser import BaseParser

class DOCXParser(BaseParser):
    supported_extensions = frozenset({".docx", ".doc"})
    parser_name = "python-docx"

    def parse(self, path: str | Path) -> Document:
        path = self.validate_path(path)
        if path.suffix.lower() == ".doc":
            raise RuntimeError("Legacy .doc parsing requires conversion to .docx before ingestion; install a converter service such as LibreOffice headless.")
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise RuntimeError("Missing dependency 'python-docx'. Install requirements.txt.") from exc
        doc = WordDocument(str(path))
        tokens: list[Token] = []
        blocks: list[TextBlock] = []
        tables: list[Table] = []
        cells: list[TableCell] = []
        y = 0.0
        for i, para in enumerate(doc.paragraphs):
            text = para.text or ""
            if not text:
                continue
            bbox = BoundingBox(0.0, y, max(100.0, len(text) * 5.0), y + 12.0)
            block_id = f"p1.block{i}"
            blocks.append(TextBlock(id=block_id, bbox=bbox, page_number=1, block_number=i, text=text))
            tokens.append(Token(id=f"{block_id}.token", text=text, bbox=bbox, page_number=1, source="ocr", block_id=block_id))
            y += 16.0
        for ti, table in enumerate(doc.tables):
            table_cell_ids: list[str] = []
            start_y = y
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    text = "\n".join(p.text for p in cell.paragraphs).strip()
                    bbox = BoundingBox(c * 120.0, y + r * 24.0, (c + 1) * 120.0, y + (r + 1) * 24.0)
                    cid = f"p1.table{ti}.r{r}.c{c}"
                    table_cell_ids.append(cid)
                    cells.append(TableCell(id=cid, bbox=bbox, page_number=1, row=r, column=c, text=text))
                    if text:
                        tokens.append(Token(id=f"{cid}.token", text=text, bbox=bbox, page_number=1, source="ocr"))
            table_bbox = BoundingBox(0.0, start_y, 120.0 * max((len(row.cells) for row in table.rows), default=1), y + 24.0 * len(table.rows))
            tables.append(Table(id=f"p1.table{ti}", bbox=table_bbox, page_number=1, cell_ids=table_cell_ids, source="geometry"))
            y = table_bbox.y1 + 16.0
        images = _docx_images(path)
        metadata = {"core_properties": {k: str(getattr(doc.core_properties, k)) for k in ("author", "title", "subject", "created", "modified")}}
        page = Page(number=1, bbox=BoundingBox(0.0, 0.0, 612.0, max(792.0, y + 24.0)), rotation=0, text_blocks=blocks, tokens=tokens, tables=tables, cells=cells, images=images)
        return Document(source_path=str(path), page_count=1, metadata=metadata, pages=[page], parser=self.parser_name, schema_version="engineering.document.v1")

def _docx_images(path: Path) -> list[Image]:
    images: list[Image] = []
    with ZipFile(path) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
    for idx, name in enumerate(media):
        images.append(Image(id=f"p1.image{idx}", bbox=BoundingBox(0, 0, 0, 0), page_number=1, xref=idx, name=name))
    return images
